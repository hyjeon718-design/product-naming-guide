from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── Brand fonts (AIWORKX 전사 표준) ──────────────────────────────────────────
FONT_KR  = "Noto Sans KR"
FONT_EN  = "Poppins"

# ── Colors ────────────────────────────────────────────────────────────────────
C_BLUE   = RGBColor(0x1C, 0x6A, 0xFF)
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_GRAY   = RGBColor(0x6B, 0x6B, 0x6B)
C_LGRAY  = RGBColor(0xF5, 0xF5, 0xF7)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_WARN   = RGBColor(0xFF, 0x8C, 0x00)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run_font(run, size=Pt(10), bold=False, color=C_DARK, italic=False):
    run.font.name  = FONT_EN
    run.font.size  = size
    run.font.bold  = bold
    run.font.color.rgb = color
    run.font.italic = italic
    # East Asian font
    rPr = run._r.get_or_add_rPr()
    for el in rPr.findall(qn('w:rFonts')):
        rPr.remove(el)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT_EN)
    rFonts.set(qn('w:hAnsi'),    FONT_EN)
    rFonts.set(qn('w:eastAsia'), FONT_KR)
    rPr.insert(0, rFonts)

def add_para(text, size=Pt(10), bold=False, color=C_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
             italic=False, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p

def add_heading(text, level=1):
    sizes   = {1: Pt(18), 2: Pt(13), 3: Pt(11)}
    befores = {1: 14,     2: 12,     3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(befores.get(level, 8))
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, Pt(11)),
                 bold=True, color=C_BLUE if level == 1 else C_DARK)
    # Bottom border for H1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1C6AFF')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_bg='1C6AFF'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        set_run_font(run, size=Pt(9), bold=True, color=C_WHITE)
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = 'F5F5F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            set_cell_bg(cells[ci], bg)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(9))
    # Column widths
    if col_widths:
        for ri2, row in enumerate(t.rows):
            for ci2, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci2])
    doc.add_paragraph()  # spacing after table

def add_note(text, color=C_WARN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    set_run_font(run, size=Pt(9), color=color, italic=True)

def add_bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, size=Pt(10))

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT HEADER
# ══════════════════════════════════════════════════════════════════════════════

# Blue top bar (simulate with shaded paragraph)
p_top = doc.add_paragraph()
p_top.paragraph_format.space_before = Pt(0)
p_top.paragraph_format.space_after  = Pt(0)
pPr = p_top._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'),   'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'),  '1C6AFF')
pPr.append(shd)
run = p_top.add_run('AIWORKX  |  Brand & Impact Team')
set_run_font(run, size=Pt(9), bold=True, color=C_WHITE)
p_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Title
add_para('Microsoft Graph API 앱 등록 요청서',
         size=Pt(22), bold=True, color=C_DARK,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=2)

add_para('Microsoft 365 연동 자동화를 위한 Azure App Registration 요청',
         size=Pt(11), color=C_GRAY, space_after=10)

# Meta info table
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Table Grid'
meta_data = [
    ('요청 부서', 'Brand & Impact Team (마케팅)'),
    ('요청일',   '2026.03.19'),
    ('요청자',   'impact.brand@aiworkx.ai'),
    ('목적',    '전사 브랜드 네이밍 협업 자동화 — 사내 협업 툴 구축'),
]
for ri, (label, val) in enumerate(meta_data):
    cells = meta_table.rows[ri].cells
    set_cell_bg(cells[0], 'E8F0FF')
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(label)
    set_run_font(r0, size=Pt(9), bold=True, color=C_BLUE)
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after  = Pt(2)
    p1 = cells[1].paragraphs[0]
    r1 = p1.add_run(val)
    set_run_font(r1, size=Pt(9), color=C_DARK)
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(2)
    cells[0].width = Cm(3.5)
    cells[1].width = Cm(12.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 요청 배경
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. 요청 배경', level=1)

add_para('마케팅팀에서 신규 제품·솔루션 네이밍 검토 업무를 자동화하기 위한 사내 협업 도구를 구축 중입니다. '
         '현재 수작업으로 진행되는 프로세스를 자동화하여 업무 효율을 높이고자 합니다.',
         space_after=8)

add_heading('현재 프로세스 (수동)', level=2)
for item in [
    '네이밍 검토 Word 문서 수동 작성',
    'SharePoint에 수동 업로드 및 링크 복사',
    'Teams 채널에 링크 수동 공유 및 피드백 요청',
    '피드백 수집 후 이메일 수동 배포',
]:
    add_bullet(item)

add_heading('자동화 목표', level=2)
for item in [
    'Claude Code 스킬 실행 → Word 검토 문서 자동 생성',
    'SharePoint 지정 폴더에 자동 업로드',
    'Teams 채널에 자동 알림 및 공동편집 링크 공유',
    '완성본 이메일 자동 배포 (전사/관련 부서)',
]:
    add_bullet(item)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. 요청 사항
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. 요청 사항', level=1)
add_para('Azure Portal에서 앱 등록(App Registration) 1건 생성을 요청드립니다.', space_after=8)

add_heading('앱 기본 정보', level=2)
add_table(
    headers=['항목', '값'],
    rows=[
        ('앱 이름',      'AIWORKX Naming Guide'),
        ('계정 유형',    '이 조직 디렉터리의 계정만 (단일 테넌트)'),
        ('리디렉션 URI', '없음 (백그라운드 서비스 — 사용자 로그인 불필요)'),
        ('앱 유형',     'Daemon / Background Service'),
    ],
    col_widths=[4.0, 12.0],
)

add_heading('필요한 API 권한 (Microsoft Graph)', level=2)
add_table(
    headers=['권한', '유형', '용도'],
    rows=[
        ('Files.ReadWrite.All',  'Application', 'SharePoint 파일 업로드·조회'),
        ('Sites.ReadWrite.All',  'Application', 'SharePoint 사이트 접근'),
        ('Mail.Send',            'Application', 'Outlook 이메일 자동 발송'),
        ('User.Read.All',        'Application', '발신자 계정 확인'),
    ],
    col_widths=[5.5, 3.5, 7.0],
)
add_note('⚠  모두 Application 권한 (위임 권한 아님) — 백그라운드 자동 실행용입니다.', color=C_WARN)

add_heading('관리자 동의 (Admin Consent)', level=2)
for item in [
    '위 권한 모두 관리자 동의(Admin Consent) 필요합니다.',
    '등록 후: Azure Portal → Enterprise Applications → 해당 앱 → Permissions → "Grant admin consent for AIWORKX" 클릭',
]:
    add_bullet(item)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 완료 후 전달 필요 정보
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. 설정 완료 후 전달 필요 정보', level=1)
add_para('앱 등록 완료 시 아래 3가지 값을 마케팅팀으로 전달 부탁드립니다.', space_after=8)

add_table(
    headers=['항목', '확인 위치'],
    rows=[
        ('Application (Client) ID', '앱 등록 → 개요'),
        ('Directory (Tenant) ID',   '앱 등록 → 개요'),
        ('Client Secret 값',         '앱 등록 → 인증서 및 암호 → 새 클라이언트 암호 생성 후 즉시 복사'),
    ],
    col_widths=[5.5, 10.5],
)
add_note('⚠  Client Secret은 생성 직후 1회만 확인 가능합니다. 반드시 즉시 복사하여 전달 부탁드립니다.', color=C_WARN)

add_heading('SharePoint 사이트 URL 확인 요청', level=2)
add_para('마케팅팀 SharePoint 사이트 URL도 함께 확인 부탁드립니다.', space_after=4)
add_bullet('경로: Teams → 마케팅 채널 → 파일 탭 → "SharePoint에서 열기" → 주소창 URL')
add_bullet('예시: https://aiworkx.sharepoint.com/sites/marketing')

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. 보안 및 관리 기준
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. 보안 및 관리 기준', level=1)

add_table(
    headers=['항목', '내용'],
    rows=[
        ('사용 범위',     '마케팅팀 전용 (Brand & Impact Team)'),
        ('접근 데이터',   'SharePoint 마케팅 폴더, 마케팅팀 발신 이메일만'),
        ('자격증명 보관', '로컬 .env 파일 (Git 미포함, .gitignore 처리)'),
        ('Secret 만료',  '1년 (만료 전 갱신 요청 예정)'),
        ('삭제 요청 시',  '요청 즉시 앱 등록 삭제 가능'),
    ],
    col_widths=[4.0, 12.0],
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. 활용 범위
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. 동일 자격증명으로 활용 가능한 기능', level=1)
add_para('이번에 등록하는 앱 하나로 아래 기능을 모두 커버합니다.', space_after=6)
for item in [
    'SharePoint 문서 자동 업로드 (네이밍 검토 Word 파일)',
    'Teams 파일 공유 링크 자동 생성 및 채널 알림',
    'Outlook 이메일 자동 발송 (마케팅 자료 배포)',
    '향후 확장: 캘린더 일정 생성, Planner 태스크 연동',
]:
    add_bullet(item)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
divider()

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run('문의: impact.brand@aiworkx.ai  |  Brand & Impact Team  |  AIWORKX  |  2026.03.19')
set_run_font(run_f, size=Pt(8), color=C_GRAY)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/n-141macbookair-3/Documents/my-first-skill/product-naming-guide/연동가이드/AIWORKX_IT_API요청서_20260319.docx'
doc.save(out)
print(f'Saved → {out}')
